from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:/Users/kong/Documents/Codex/2026-05-11/files-mentioned-by-the-user-prd")
TEMPLATE_PATH = ROOT / "template.docx"
ASCII_OUTPUT_PATH = ROOT / "feasibility_report_pv.docx"
FINAL_OUTPUT_PATH = ROOT / "辽宁某产业园20MW屋顶分布式光伏项目可行性分析报告.docx"


TITLE_LINES = [
    "辽宁某产业园20MW屋顶",
    "分布式光伏发电项目",
    "可行性分析报告",
]


SECTIONS = {
    "总论": [
        (
            "label",
            "项目背景",
            "这次课程报告我选择的是产业园屋顶分布式光伏项目。之所以选这个题目，主要是因为近几年“双碳”目标持续推进，"
            "很多制造业企业都开始关注用电成本和绿色用能问题。辽宁是东北地区重要的工业基地，装备制造、冶金、石化等产业基础较强，"
            "不少园区白天用电需求比较稳定，分布式光伏项目也越来越受到重视，"
            "所以把它作为可行性分析对象，既贴近现实，也比较符合工程经济课程的学习内容。",
        ),
        (
            "label",
            "建设内容",
            "本项目拟利用辽宁某产业园已经建成的厂房、仓库和配套建筑屋面，建设总装机容量20MWp的屋顶分布式光伏发电系统。"
            "项目采用“自发自用、余电上网”的运行方式，同时配套逆变器、监控系统、消防设施、并网装置和后期运维系统。"
            "按照初步设想，建设期约为7个月，运营期按25年考虑。",
        ),
        (
            "label",
            "结论摘要",
            "从屋面条件、园区负荷、并网消纳、环境影响和经济测算几个方面来看，这个项目都具备一定的实施基础。"
            "虽然本文的测算仍然是基于公开资料和合理假设完成的，但整体结果说明该项目具有较好的可行性，"
            "在实际中也有一定参考价值。",
        ),
    ],
    "需求预测和拟建规模": [
        (
            "label",
            "需求分析",
            "假设该产业园内的企业主要以装备制造、金属加工、仓储物流等产业为主，这类企业通常白天用电较集中，"
            "而且全年生产比较稳定。若园区年用电量约为3,500万kWh，那么白天的基础负荷就能够为光伏消纳提供条件。"
            "光伏发出来的电优先在园区内部消耗，一方面可以减少外购电费用，另一方面也有助于提高企业的绿色用能比例。",
        ),
        (
            "label",
            "规模论证",
            "根据屋面面积、遮挡情况、消防通道以及检修空间等条件进行估算，园区大约可以提供15.5万平方米左右的可利用屋面。"
            "按照目前较常见的组件布置方式测算，可形成约20MWp的装机容量。若继续扩大规模，可能会受到屋面承载、"
            "配电容量和自用比例下降等因素限制，因此20MWp可以认为是一个比较合适的建设规模。",
        ),
        (
            "label",
            "建设目标",
            "本文按首年发电量2,050万kWh进行测算，其中约82%的电量由园区内部直接消纳，其余部分上网。"
            "这样既能满足园区白天的部分用电需求，也能兼顾项目收益和并网运行的稳定性。",
        ),
    ],
    "自然资源、原材料、燃料及公用设施情况": [
        (
            "label",
            "自然资源条件",
            "辽宁的太阳能资源条件虽然比不上西北地区，但在东北地区中属于较适合开发的一类，整体年辐照情况比较稳定，适合建设分布式光伏项目。"
            "结合近年的气象资料和东北、环渤海地区屋顶光伏项目的一般情况，本文采用相对保守的年利用小时数进行估算，"
            "这样得出的结果会更稳妥一些。",
        ),
        (
            "label",
            "材料与设备供应",
            "项目需要的主要设备包括N型单晶组件、组串式逆变器、支架、线缆、并网柜和监控平台等。"
            "这些产品在国内已经比较成熟，辽宁及环渤海地区也有较多设备厂家和施工单位，"
            "因此从采购和供货角度看，项目实施难度不算大。",
        ),
        (
            "label",
            "燃料与公用工程",
            "光伏项目在运行阶段不需要消耗煤、油、气等传统燃料，日常主要是少量清洗用水、值守用电和网络通信资源。"
            "如果园区原有10kV配电设施、道路和消防等基础条件较完善，那么施工、并网和后续运维都比较容易开展。",
        ),
    ],
    "建厂条件和方案": [
        (
            "label",
            "建设条件",
            "项目场址位于已经投运的产业园区内，土地和房屋权属假定比较清晰，园区外部交通条件也较成熟。"
            "施工时可以利用现有道路和装卸区域进行材料运输与设备吊装。屋面类型以彩钢瓦和部分混凝土屋面为主，"
            "总体上适合建设屋顶光伏，但正式实施前仍然要逐栋复核荷载、防水和结构安全情况。",
        ),
        (
            "label",
            "方案比选",
            "这里主要比较两种思路。第一种是“自发自用、余电上网”，优点是能更好地利用园区白天负荷，减少企业购电支出；"
            "第二种是“全额上网”，虽然运行方式更简单，但收益会更加依赖上网电价。"
            "综合来看，工商业屋顶光伏更适合采用第一种方式，因此本文建议选择“自发自用、余电上网”方案。",
        ),
        (
            "label",
            "推荐方案",
            "从建设和后期管理角度看，推荐采用统一投资、集中运维、分区接入的方式。"
            "彩钢瓦屋面可优先采用专用夹具固定，尽量减少穿孔带来的渗漏风险；混凝土屋面则可以采用压载或锚栓等方式。"
            "同时还要留出检修通道和消防间距，保证后期运行安全。",
        ),
    ],
    "设计方案": [
        (
            "label",
            "系统配置",
            "项目拟采用620Wp级N型单晶硅组件，并配置1500V直流系统和320kW级组串式逆变器。"
            "另外还要设置必要的直流隔离、交流保护、汇流监测和远程运维平台。"
            "总体设计思路是分散布置、就近接入、统一监控，这样既方便发电，也便于后期检修。",
        ),
        (
            "label",
            "电气与结构设计",
            "在布置组件时，要尽量避开女儿墙、采光带和通风设备，同时合理控制行距和检修通道。"
            "电气部分要考虑防雷接地、绝缘保护和消防联动等要求；结构部分则要重点考虑风荷载、屋面连接方式、"
            "防腐和排水问题，尽量不要影响原有建筑的正常使用。",
        ),
        (
            "label",
            "智慧运维设计",
            "为了提高运行效率，项目还应配置发电监控、故障告警和基础报表功能。"
            "如果条件允许，也可以预留后续接入储能或综合能源管理系统的接口，"
            "这样将来在园区能源优化方面会更灵活。",
        ),
    ],
    "环境保护": [
        (
            "label",
            "建设期环保措施",
            "施工阶段的环境影响主要体现在设备运输、施工噪声、包装废弃物和少量扬尘等方面。"
            "这些问题总体可控，可以通过合理安排施工时间、及时清运废弃物、做好屋面作业防护等措施来减少影响，"
            "并尽量避开园区企业生产较忙的时段。",
        ),
        (
            "label",
            "运营期环境影响",
            "项目运行后基本不会产生废气和废水，也不会像传统火电那样排放燃烧污染物，因此环境压力相对较小。"
            "不过仍然要注意逆变器噪声、报废元件回收和组件寿命结束后的处理问题。",
        ),
        (
            "label",
            "综合效益",
            "按首年发电量2,050万kWh估算，项目每年可减少二氧化碳排放约1.17万吨。"
            "除了减排作用外，它还能在一定程度上减少常规能源消耗，对提升园区绿色形象也有帮助。",
        ),
    ],
    "企业组织、劳动定员和人员培训": [
        (
            "label",
            "组织模式",
            "从组织方式看，建设期可以由业主设立专门项目组，分别负责工程、采购、安全和协调工作；"
            "运营期则可以交由园区能源管理部门或专业运维单位统一管理。",
        ),
        (
            "label",
            "劳动定员",
            "项目投运后不需要太多常驻人员，初步考虑配置6名核心运维人员即可，"
            "包括项目负责人、电气技术人员、现场运维人员和安全管理人员。"
            "像组件清洗和专项检修这类工作，也可以根据需要外包处理。",
        ),
        (
            "label",
            "培训要求",
            "培训内容应包括高处作业、电气安全、消防应急、监控平台使用和日常缺陷处理等。"
            "关键岗位最好做到持证上岗，并定期复训，这样更有利于保障项目长期稳定运行。",
        ),
    ],
    "实施进度的建议": [
        (
            "label",
            "总体安排",
            "考虑到项目位于正常运行中的产业园内，实施时应按照“先审批、后施工、分批并网”的思路推进。"
            "其中比较关键的环节是屋面复核、设备采购和并网手续办理，这几个部分如果衔接不好，容易影响总工期。",
        ),
        ("caption", "表1  项目实施进度建议"),
        (
            "table",
            {
                "headers": ["阶段", "主要工作", "建议工期"],
                "rows": [
                    ["项目立项与可研", "资料收集、方案论证、投资测算、项目决策", "0.5个月"],
                    ["初设与并网报装", "结构复核、系统设计、接入审查、手续申报", "1.0个月"],
                    ["招标采购", "组件、逆变器、支架及施工单位招采", "1.0个月"],
                    ["屋面处理与基础施工", "加固、防水节点处理、支架安装", "1.5个月"],
                    ["设备安装与调试", "组件敷设、电气接线、监控接入、单体调试", "2.0个月"],
                    ["并网验收与试运行", "联合验收、并网送电、消缺与移交", "1.0个月"],
                ],
                "widths": [Cm(3.7), Cm(9.6), Cm(3.2)],
                "alignments": ["center", "left", "center"],
            },
        ),
        (
            "plain",
            "按照上表安排，项目总工期大约为7个月。如果并网审批比较顺利，且天气影响不大，工期还有适当压缩的可能。",
        ),
    ],
    "投资估算和资金筹措": [
        (
            "label",
            "投资估算",
            "本文按单位静态投资约3.6元/W进行估算，项目总投资约为7,200万元。"
            "该数值已经把设备采购、安装施工、并网接入、设计监理和预备费等因素考虑进去，"
            "但因为属于课程报告中的初步测算，所以和实际工程中的最终投资金额仍可能存在差异。",
        ),
        ("caption", "表2  项目投资估算表"),
        (
            "table",
            {
                "headers": ["费用项目", "金额（万元）", "占比"],
                "rows": [
                    ["光伏组件及逆变器", "4,200", "58.3%"],
                    ["支架、线缆及配电设备", "1,200", "16.7%"],
                    ["安装工程费", "900", "12.5%"],
                    ["设计、监理及手续费用", "300", "4.2%"],
                    ["基本预备费及其他", "600", "8.3%"],
                    ["合计", "7,200", "100.0%"],
                ],
                "widths": [Cm(8.2), Cm(4.0), Cm(4.3)],
                "alignments": ["left", "center", "center"],
            },
        ),
        (
            "label",
            "资金筹措",
            "资金方面可以采用“资本金+银行贷款”的方式筹措。"
            "例如资本金按30%考虑，大约需要2,160万元；其余70%通过贷款解决，大约为5,040万元。"
            "如果业主希望减轻一次性投资压力，后续也可以再考虑融资租赁等方式。",
        ),
    ],
    "财务分析、国民经济分析和社会分析": [
        (
            "label",
            "测算假设",
            "财务测算中，本文假设首年发电量为2,050万kWh，自发自用比例为82%，自用电价为0.70元/kWh，"
            "余电上网电价为0.38元/kWh，年衰减率按0.5%考虑，初始年运维成本按230万元估算。"
            "这些参数主要用于课程分析，后续如果进入实际项目阶段，还需要结合合同电价和融资条件进一步修正。",
        ),
        ("caption", "表3  主要经济与综合效益指标"),
        (
            "table",
            {
                "headers": ["指标", "测算值"],
                "rows": [
                    ["首年营业收入", "约1,317万元"],
                    ["年均运维成本（初始年）", "约230万元"],
                    ["静态投资回收期", "约6.8年"],
                    ["税前财务内部收益率", "约13.5%"],
                    ["按8%折现的净现值", "约3,382万元"],
                    ["年减排二氧化碳", "约1.17万吨"],
                ],
                "widths": [Cm(8.2), Cm(8.3)],
                "alignments": ["left", "center"],
            },
        ),
        (
            "label",
            "综合评价",
            "从财务结果看，项目具有较稳定的收益来源，投资回收期也处在可以接受的范围内；"
            "从社会和环境效益看，它能够减少碳排放、改善园区用能结构，也有利于提升企业绿色发展形象。"
            "综合以上分析，本文认为该项目总体上是可行的，可以作为后续进一步深化研究和实际论证的基础。",
        ),
    ],
    "参考文献": [
        ("reference", "[1] 国家能源局. 2024年上半年光伏发电建设情况[EB/OL]. 2024-07-25."),
        ("reference", "[2] 中国政府网. 《分布式光伏发电开发建设管理办法》政策解读[EB/OL]. 2025-01-23."),
        ("reference", "[3] 中国气象局. 2024年中国风能太阳能资源年景公报[EB/OL]. 2025-01-24."),
        ("reference", "[4] 辽宁省人民政府办公厅. 辽宁省人民政府办公厅关于印发辽宁省“十四五”能源发展规划的通知[EB/OL]. 2023-03-13."),
        ("reference", "[5] 辽宁省人民政府. 关于印发《辽宁省分布式光伏发电开发建设管理实施细则》的通知[EB/OL]. 2026-02-06."),
        ("reference", "[6] 辽宁省人民政府. 绿电领跑全省 辽宁新能源装机规模超越火电成为第一大电源类型[EB/OL]. 2026-04-21."),
    ],
}


def remove_between(start_para: Paragraph, end_para: Paragraph) -> None:
    current = start_para._p.getnext()
    while current is not None and current is not end_para._p:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def insert_paragraph_after(anchor: Paragraph | Table, text: str = "", style: str | None = None) -> Paragraph:
    anchor_el = anchor._p if isinstance(anchor, Paragraph) else anchor._tbl
    parent = anchor._parent
    new_p = OxmlElement("w:p")
    anchor_el.addnext(new_p)
    para = Paragraph(new_p, parent)
    if style:
        para.style = style
    if text:
        para.add_run(text)
    return para


def insert_table_after(anchor: Paragraph | Table, doc: Document, rows: int, cols: int) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    anchor_el = anchor._p if isinstance(anchor, Paragraph) else anchor._tbl
    anchor_el.addnext(table._tbl)
    return table


def set_run_font(
    run,
    size_pt: float = 12,
    *,
    bold: bool | None = None,
    font_name: str = "Microsoft YaHei",
    east_asia: str = "微软雅黑",
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def format_body_paragraph(para: Paragraph, *, first_line: bool = True) -> None:
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.5
    for run in para.runs:
        set_run_font(run, 12)


def format_reference_paragraph(para: Paragraph) -> None:
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = para.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.35
    for run in para.runs:
        set_run_font(run, 11.5)


def set_style_fonts(doc: Document) -> None:
    for style_name, size, bold in [("Normal", 12, False), ("Heading 1", 16, True)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style._element.rPr.rFonts.set(qn("w:cs"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold


def make_labeled_paragraph(anchor: Paragraph | Table, label: str, body: str) -> Paragraph:
    para = insert_paragraph_after(anchor)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = para.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.5
    label_run = para.add_run(f"{label}：")
    set_run_font(label_run, 12, bold=True)
    body_run = para.add_run(body)
    set_run_font(body_run, 12)
    return para


def make_caption(anchor: Paragraph | Table, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, text=text)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = para.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.2
    for run in para.runs:
        set_run_font(run, 11, bold=True)
    return para


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        child = tc_mar.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def build_table(anchor: Paragraph | Table, doc: Document, spec: dict) -> Table:
    rows = spec["rows"]
    headers = spec["headers"]
    widths = spec["widths"]
    aligns = spec["alignments"]
    table = insert_table_after(anchor, doc, len(rows) + 1, len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    for i, head in enumerate(headers):
        cell = header_cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        para.text = head
        for run in para.runs:
            set_run_font(run, 10.5, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            align_key = aligns[c_idx]
            if align_key == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_key == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.text = text
            for run in para.runs:
                set_run_font(run, 10.5)
    return table


def style_heading(para: Paragraph, *, new_page: bool = True) -> None:
    para.style = "Heading 1"
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = para.paragraph_format
    fmt.page_break_before = new_page
    fmt.keep_with_next = True
    fmt.space_after = Pt(8)
    for run in para.runs:
        set_run_font(run, 16, bold=True)


def replace_cover(doc: Document) -> None:
    title_para = doc.paragraphs[4]
    title_para.text = ""
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing = 1.25
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(TITLE_LINES):
        run = title_para.add_run(line)
        if idx < len(TITLE_LINES) - 1:
            run.add_break(WD_BREAK.LINE)
        set_run_font(run, 24)

    info_values = {
        18: "课程名称：工程经济导论",
        19: "班级：____________________",
        20: "姓名：____________________",
        21: "学号：____________________",
    }
    for idx, text in info_values.items():
        para = doc.paragraphs[idx]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Cm(2)
        for run in para.runs:
            set_run_font(run, 18)


def trim_sections(doc: Document, heading_map: dict[str, Paragraph]) -> None:
    ordered = [heading_map[name] for name in SECTIONS]
    for current, nxt in zip(ordered, ordered[1:]):
        remove_between(current, nxt)


def populate_sections(doc: Document, heading_map: dict[str, Paragraph]) -> None:
    for idx, (heading_text, blocks) in enumerate(SECTIONS.items()):
        heading = heading_map[heading_text]
        style_heading(heading, new_page=True)
        anchor: Paragraph | Table = heading
        for block in blocks:
            kind = block[0]
            if kind == "label":
                anchor = make_labeled_paragraph(anchor, block[1], block[2])
            elif kind == "plain":
                anchor = insert_paragraph_after(anchor, block[1])
                format_body_paragraph(anchor, first_line=True)
            elif kind == "caption":
                anchor = make_caption(anchor, block[1])
            elif kind == "table":
                anchor = build_table(anchor, doc, block[1])
            elif kind == "reference":
                anchor = insert_paragraph_after(anchor, block[1])
                format_reference_paragraph(anchor)
        if idx == 0:
            heading.paragraph_format.page_break_before = True


def find_heading_map(doc: Document) -> dict[str, Paragraph]:
    mapping: dict[str, Paragraph] = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in SECTIONS:
            mapping[text] = para
    missing = [name for name in SECTIONS if name not in mapping]
    if missing:
        raise ValueError(f"Missing headings in template: {missing}")
    return mapping


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))
    set_style_fonts(doc)
    replace_cover(doc)
    heading_map = find_heading_map(doc)
    trim_sections(doc, heading_map)
    populate_sections(doc, heading_map)
    doc.save(str(ASCII_OUTPUT_PATH))
    doc.save(str(FINAL_OUTPUT_PATH))
    print(f"Saved: {ASCII_OUTPUT_PATH}")
    print(f"Saved: {FINAL_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
